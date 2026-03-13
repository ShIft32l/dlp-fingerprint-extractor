import os
from pathlib import Path
from parser import extract_document_text
from docx import Document
from reportlab.pdfgen import canvas

def create_dummy_files(samples_dir: str):
    os.makedirs(samples_dir, exist_ok=True)
    
    # 1. TXT
    txt_path = Path(samples_dir) / "dummy.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("This is a dummy TXT file.\nLine 2 of text.")
        
    # 2. DOCX
    docx_path = Path(samples_dir) / "dummy.docx"
    doc = Document()
    doc.add_heading('Dummy Document', 0)
    doc.add_paragraph('This is a dummy DOCX file.')
    doc.add_paragraph('Special characters: 12345 !@#$%^&*()')
    doc.save(str(docx_path))
    
    # 3. PDF
    pdf_path = Path(samples_dir) / "dummy.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 800, "This is a dummy PDF file.")
    c.drawString(100, 780, "Extracting text block from PDF.")
    c.save()

    return txt_path, docx_path, pdf_path

def run_tests():
    # Append start to session log
    brain_dir = Path("../.brain")
    brain_dir.mkdir(exist_ok=True)
    session_log = brain_dir / "session_log.txt"
    
    def log_append(message):
        with open(session_log, "a", encoding="utf-8") as f:
            f.write(message + "\n")
            
    log_append("[auto] START phase-02-parser")
    
    samples_dir = "../samples"
    txt, docx, pdf = create_dummy_files(samples_dir)
    print("Created dummy files.")
    log_append("[auto] DONE task: Create test files (dummy.txt, docx, pdf)")
    
    # Test txt
    text = extract_document_text(txt)
    print(f"TXT Content:\n{text}")
    assert "dummy TXT" in text, "TXT extraction failed"
    
    # Test docx
    text = extract_document_text(docx)
    print(f"DOCX Content:\n{text}")
    assert "dummy DOCX" in text, "DOCX extraction failed"
    
    # Test pdf
    text = extract_document_text(pdf)
    print(f"PDF Content:\n{text}")
    assert "dummy PDF" in text, "PDF extraction failed"
    
    print("All tests passed.")
    log_append("[auto] DONE task: Run tests and verify extraction")
    log_append("[auto] END phase-02-parser ✅")

if __name__ == "__main__":
    run_tests()
